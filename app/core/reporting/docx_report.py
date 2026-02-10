from pathlib import Path
from typing import Any

from docx import Document


def build_docx_report(report_data: dict[str, Any], out_path: Path) -> None:
    doc = Document()
    doc.add_heading("Forensic Linguistics Desktop v1.0", level=1)
    doc.add_paragraph("Внимание: программа не устанавливает автора и смысл автоматически; требуется экспертная интерпретация.")

    doc.add_heading("Методика и воспроизводимость", level=2)
    reproducibility = report_data.get("reproducibility", {})
    for k, v in reproducibility.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Режим 1. Статистический профиль", level=2)
    for k, v in report_data.get("mode1", {}).items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Режим 2. Контекстный анализ", level=2)
    for row in report_data.get("mode2_fragments", []):
        doc.add_paragraph(f"{row}")

    doc.add_heading("Режим 3. Сравнение профилей", level=2)
    for k, v in report_data.get("mode3", {}).items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Ограничения", level=2)
    for w in report_data.get("warnings", []):
        doc.add_paragraph(f"• {w}")

    doc.save(out_path)
