import hashlib
import json
import platform
import threading
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
import tkinter as tk

from app.core.authorship.advanced_analysis import compare_texts_detailed
from app.core.backends.natasha_backend import NatashaBackend
from app.core.backends.pymorphy_backend import PymorphyBackend
from app.core.context.hints import contextual_hints
from app.core.metrics.lex import lex_metrics
from app.core.metrics.ngrams import char_ngrams
from app.core.metrics.orth import orth_metrics
from app.core.metrics.pos import pos_ngrams, pos_profile
from app.core.metrics.punct import punct_metrics
from app.core.metrics.quality import quality_flags
from app.core.preprocess.normalize import normalize_text
from app.core.reporting.docx_report import build_docx_report
from app.core.reporting.exporters import export_json, export_tokens_csv
from app.storage.db import ProjectStore
from app.ui.tabs.authorship_tab import AuthorshipTab
from app.ui.tabs.context_tab import ContextTab
from app.ui.tabs.report_tab import ReportTab
from app.ui.tabs.statistical_tab import StatisticalTab


class MainWindow(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("Forensic Linguistics Desktop v1.0")
        self.geometry("1300x900")
        self.cancel_event = threading.Event()

        self.project_dir = Path.cwd() / "project"
        self.store = ProjectStore(self.project_dir)
        self.case_id = self.store.create_case("Дело по умолчанию")

        self.backends = {
            "natasha": NatashaBackend(),
            "pymorphy3": PymorphyBackend(),
        }

        self._build_ui()
        self._bind_global_hotkeys()

    def _bind_global_hotkeys(self) -> None:
        self.bind_all("<Control-v>", self._paste_focused_widget)
        self.bind_all("<Control-V>", self._paste_focused_widget)
        self.bind_all("<Shift-Insert>", self._paste_focused_widget)

    def _paste_focused_widget(self, _: tk.Event) -> str:
        widget = self.focus_get()
        if widget is None:
            return "break"
        try:
            widget.event_generate("<<Paste>>")
        except Exception:
            pass
        return "break"

    def _build_ui(self) -> None:
        header = ttk.Frame(self)
        header.pack(fill="x", padx=8, pady=6)
        ttk.Button(header, text="Открыть txt/docx", command=self._import_text).pack(side="left")
        ttk.Button(header, text="Диагностика окружения", command=self._show_diagnostics).pack(side="left", padx=6)
        self.backend_var = tk.StringVar(value="natasha")
        ttk.Label(header, text="Бэкенд:").pack(side="left", padx=(12, 4))
        ttk.Combobox(header, state="readonly", values=["natasha", "pymorphy3"], textvariable=self.backend_var, width=12).pack(side="left")

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True)

        self.stat_tab = StatisticalTab(self.notebook, self)
        self.ctx_tab = ContextTab(self.notebook, self)
        self.auth_tab = AuthorshipTab(self.notebook, self)
        self.rep_tab = ReportTab(self.notebook, self)

        self.notebook.add(self.stat_tab, text="1) Статистический профиль")
        self.notebook.add(self.ctx_tab, text="2) Контекстный анализ")
        self.notebook.add(self.auth_tab, text="3) Автоворедение")
        self.notebook.add(self.rep_tab, text="4) Мастер отчёта")

    def _show_diagnostics(self) -> None:
        info = {
            "python": platform.python_version(),
            "implementation": platform.python_implementation(),
            "platform": platform.platform(),
            "backend_selected": self.backend_var.get(),
        }
        messagebox.showinfo("Диагностика", json.dumps(info, ensure_ascii=False, indent=2))

    def _import_text(self) -> None:
        path = filedialog.askopenfilename(filetypes=[("Text", "*.txt *.docx")])
        if not path:
            return
        p = Path(path)
        if p.suffix.lower() == ".docx":
            from docx import Document

            doc = Document(path)
            text = "\n".join(par.text for par in doc.paragraphs)
        else:
            text = p.read_text(encoding="utf-8", errors="ignore")
        self.stat_tab.set_text(text)

    def run_analysis_background(self, mode: str, text: str, callback) -> None:
        self.cancel_event.clear()

        def worker() -> None:
            try:
                norm_text = normalize_text(text)
                backend_name = self.backend_var.get()
                result = self.backends[backend_name].analyze(norm_text)
                warnings = quality_flags(result)
                manifest = {
                    "sha256": hashlib.sha256(norm_text.encode("utf-8")).hexdigest(),
                    "backend": backend_name,
                    "python": platform.python_version(),
                    "warnings": warnings,
                    "mode": mode,
                }
                metrics = {
                    "pos_profile": pos_profile(result.tokens),
                    "pos_ngrams_2": pos_ngrams(result.tokens, 2),
                    "pos_ngrams_3": pos_ngrams(result.tokens, 3),
                    "lex": lex_metrics(result.tokens),
                    "punct": punct_metrics(norm_text),
                    "orth": orth_metrics(norm_text),
                    "char_ngrams": char_ngrams(norm_text),
                    "context_hints": contextual_hints(norm_text),
                }
                payload = {
                    "manifest": manifest,
                    "metrics": metrics,
                    "tokens": [t.__dict__ for t in result.tokens],
                    "sentences": [s.__dict__ for s in result.sentences],
                }
                out_json = self.store.results_dir / f"run_{hashlib.md5(norm_text.encode()).hexdigest()}.json"
                export_json(payload, out_json)
                doc_id = self.store.add_document(self.case_id, "исследуемый", f"Документ {mode}", "", manifest["sha256"])
                self.store.add_run(doc_id, mode, {}, backend_name, {"python": platform.python_version()}, warnings, str(out_json))
                self.after(0, lambda: callback(payload, None))
            except Exception as exc:  # noqa: BLE001
                self.after(0, lambda: callback(None, str(exc)))

        threading.Thread(target=worker, daemon=True).start()

    def export_all(self, payload: dict, target_dir: Path) -> None:
        target_dir.mkdir(parents=True, exist_ok=True)
        export_json(payload, target_dir / "report.json")
        export_tokens_csv(payload.get("tokens", []), target_dir / "tokens.csv")
        report_data = {
            "reproducibility": payload.get("manifest", {}),
            "mode1": payload.get("metrics", {}),
            "mode2_fragments": self.ctx_tab.get_fragments_text(),
            "mode3": self.auth_tab.get_comparison_summary(),
            "warnings": payload.get("manifest", {}).get("warnings", []),
        }
        build_docx_report(report_data, target_dir / "report.docx")

    def compare_with_sample(self, text_a: str, text_b: str) -> dict:
        return compare_texts_detailed(text_a, text_b)


def run_app() -> None:
    app = MainWindow()
    app.mainloop()
